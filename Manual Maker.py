import os
import time
from tkinter import filedialog
from PIL import ImageGrab, Image, ImageTk
from fpdf import FPDF
import customtkinter as ctk
from pathlib import Path
from pynput import keyboard
from fpdf.enums import XPos, YPos
from customtkinter import CTkImage
from docx import Document
from docx.shared import Inches
import mss
from PIL import Image
import win32gui
import win32ui
import win32con
import win32api
from pynput.mouse import Listener as MouseListener
from threading import Thread
from PIL import Image, ImageTk
from PIL import ImageSequence
import random

class ScreenshotApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Screenshot Documentation Tool")
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.root.resizable(False, False)
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = 400
        window_height = 300
        position_x = int(((screen_width // 2) - (window_width // 2)) * 2)
        position_y = int(((screen_height // 2) - (window_height // 2)) * 2)
        self.root.geometry(f"{window_width}x{window_height}+{position_x}+{position_y}")

        # Initialize variables
        self.is_capturing = False
        self.screenshots = []
        self.selected_screenshots = []
        self.listener = None
        self.current_keys = set()  # Track currently pressed keys

        # Configure CustomTkinter appearance
        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")

        # UI Elements
        self.start_button = ctk.CTkButton(root, text="Start Capturing", command=self.start_capturing)
        self.start_button.pack(pady=20)
        
        # Create a frame for checkboxes
        self.checkbox_frame = ctk.CTkFrame(root, fg_color="transparent")
        self.checkbox_frame.place(x=280, y=10)  # Adjust x/y to position as needed

        self.ctrl_shift_var = ctk.BooleanVar()
        self.click_var = ctk.BooleanVar()

        self.ctrl_shift_checkbox = ctk.CTkCheckBox(
            self.checkbox_frame,
            text="CTRL/SHIFT",
            variable=self.ctrl_shift_var,
            font=("Helvetica", 8),
            checkbox_width=16,
            checkbox_height=16
        )
        self.ctrl_shift_checkbox.grid(row=0, column=0, padx=(0, 4), sticky="w")

        self.click_checkbox = ctk.CTkCheckBox(
            self.checkbox_frame,
            text="CTRL/CLICK",
            variable=self.click_var,
            font=("Helvetica", 8),
            checkbox_width=16,
            checkbox_height=16,
        )
        self.click_checkbox.grid(row=1, column=0, padx=(0, 0), sticky="w")


        self.stop_button = ctk.CTkButton(root, text="Stop Capturing", command=self.stop_capturing, state="disabled")
        self.stop_button.pack(pady=20)

        self.preview_button = ctk.CTkButton(root, text="Preview Screenshots", command=self.preview_screenshots, state="disabled")
        self.preview_button.pack(pady=20)

        self.export_button = ctk.CTkButton(root, text="Export Screenshots", command=self.export_screenshots, state="disabled")
        self.export_button.pack(pady=20)

        self.status_label = ctk.CTkLabel(root, text="Status: Idle", text_color="gray", font=("Helvetica", 13), anchor="w")
        self.status_label.pack(side="bottom", fill="x", padx=10, pady=6)
        
        self.gif_frames = []
        self.gif_index = 0
        self.gif_label = ctk.CTkLabel(root, text="")
        self.gif_label.place(x=10, y=220)  # Adjust position as needed

        self.gif_x = 10  # Initial x-coordinate
        self.gif_y = 175  # Initial y-coordinate
        self.gif_step = 5  # Step size for horizontal movement
        self.gif_label.place(x=self.gif_x, y=self.gif_y)  # Initial position
        self.gif_label.lower()
        self.load_gif(r"C:\Users\NewUsername\Documents\cat-8915_128.gif")  # Path to your gif
        self.animate_gif()
        
    def animate_gif(self):
        if not self.gif_frames:
            return

        # Update the GIF frame
        frame = self.gif_frames[self.gif_index]
        self.gif_label.configure(image=frame)
        self.gif_label.image = frame

        # Update position for walking animation
        self.gif_x += self.gif_step  # Move horizontally
        print(f"Updated x-coordinate: {self.gif_x}")  # Debugging output

        # Check window width dynamically
        self.root.update_idletasks()  # Ensure dimensions are up-to-date
        window_width = self.root.winfo_width()
        print(f"Window width: {window_width}")  # Debugging output

        if self.gif_x > window_width:  # If it goes off the right edge
            self.gif_x = -self.gif_label.winfo_width()  # Restart from the left

            # Change the y-coordinate to a random position within the window height
            window_height = self.root.winfo_height()
            self.gif_y = random.randint(50, window_height - self.gif_label.winfo_height() - 50)
            print(f"Reset x-coordinate: {self.gif_x}, Updated y-coordinate: {self.gif_y}")  # Debugging output

        # Update label position
        self.gif_label.place(x=self.gif_x, y=self.gif_y)  # Move the label

        # Move to the next frame of the GIF
        self.gif_index = (self.gif_index + 1) % len(self.gif_frames)

        # Schedule the next frame
        self.root.after(100, self.animate_gif)  # Adjust timing (ms) for speed	    

    def start_capturing(self):
        self.is_capturing = True
        self.start_button.configure(state="disabled")
        self.stop_button.configure(state="normal")
        self.preview_button.configure(state="disabled")
        self.export_button.configure(state="disabled")
        self.status_label.configure(text="Status: Capturing...")
        self.start_keyboard_listener()
        self.start_mouse_listener()

    def stop_capturing(self):
        self.is_capturing = False
        self.start_button.configure(state="normal")
        self.stop_button.configure(state="disabled")
        self.preview_button.configure(state="normal")
        self.export_button.configure(state="normal")
        self.status_label.configure(text="Status: Stopped")
        self.stop_keyboard_listener()
        
        if hasattr(self, 'mouse_listener') and self.mouse_listener:
            self.mouse_listener.stop()
            self.mouse_listener = None


    def start_keyboard_listener(self):
        # Directory to save screenshots
        self.screenshots_dir = Path.home() / "Documents" / "Screenshots"
        if not self.screenshots_dir.exists():
            os.makedirs(self.screenshots_dir)

        # Keyboard key press event handlers
        def on_press(key):
            if self.is_capturing:
                try:
                    # Add the key to the set of currently pressed keys
                    self.current_keys.add(key)

                    # Check if both CTRL and SHIFT are pressed
                    if self.ctrl_shift_var.get():
                        if (keyboard.Key.ctrl_l in self.current_keys or keyboard.Key.ctrl_r in self.current_keys) and keyboard.Key.shift in self.current_keys:
                            # Take a screenshot
                            # Take a screenshot
                            timestamp = time.strftime("%Y%m%d_%H%M%S")
                            screenshot_path = self.screenshots_dir / f"step_{timestamp}.png"
                            self.capture_screenshot_with_cursor(screenshot_path)                        
                            self.screenshots.append(screenshot_path)
                            self.status_label.configure(text=f"Screenshot taken: {screenshot_path.name}", text_color="green")

                except Exception as e:
                    self.status_label.configure(text=f"Error: {str(e)}", text_color="red")

        def on_release(key):
            # Remove the key from the set of currently pressed keys
            if key in self.current_keys:
                self.current_keys.remove(key)

        # Start the listener in a non-blocking way
        self.listener = keyboard.Listener(on_press=on_press, on_release=on_release)
        self.listener.start()

    def stop_keyboard_listener(self):
        if self.listener:
            self.listener.stop()
            self.listener = None
            
    def start_mouse_listener(self):
        def on_click(x, y, button, pressed):
            if pressed and self.is_capturing:
                # Only capture if "CLICK" checkbox is checked and Ctrl key is down
                if self.is_capturing and self.click_var.get() and (win32api.GetAsyncKeyState(win32con.VK_CONTROL) & 0x8000):
                    timestamp = time.strftime("%Y%m%d_%H%M%S")
                    screenshot_path = self.screenshots_dir / f"click_{timestamp}.png"
                    self.capture_screenshot_with_cursor(screenshot_path)
                    self.screenshots.append(screenshot_path)
                    self.status_label.configure(text=f"Screenshot taken: {screenshot_path.name}", text_color="green")


        self.mouse_listener = MouseListener(on_click=on_click)
        
        # Start in its own thread to avoid GC or UI blocking
        self.mouse_thread = Thread(target=self.mouse_listener.run, daemon=True)
        self.mouse_thread.start()    
            
    def capture_screenshot_with_cursor(self, save_path):
        # Take a screenshot with MSS
        with mss.mss() as sct:
            monitor = sct.monitors[1]
            sct_img = sct.grab(monitor)
            img = Image.frombytes("RGB", sct_img.size, sct_img.rgb)

        # Get cursor info
        cursor_info = win32gui.GetCursorInfo()
        cursor_flags, cursor_handle, (cursor_x, cursor_y) = cursor_info

        if cursor_flags == win32con.CURSOR_SHOWING and cursor_handle:
            try:
                icon_info = win32gui.GetIconInfo(cursor_handle)
                hotspot_x = icon_info[1]
                hotspot_y = icon_info[2]

                # Create a device context and bitmap for the icon
                hdc_screen = win32gui.GetDC(0)
                hdc_mem = win32ui.CreateDCFromHandle(hdc_screen).CreateCompatibleDC()
                hbitmap = win32ui.CreateBitmapFromHandle(icon_info[4])  # hbmColor

                hdc_mem.SelectObject(hbitmap)

                # Extract raw image data
                bmpinfo = hbitmap.GetInfo()
                bmpstr = hbitmap.GetBitmapBits(True)

                # Create cursor image
                if bmpinfo['bmBitsPixel'] == 32:
                    cursor_img = Image.frombuffer('RGBA', (bmpinfo['bmWidth'], bmpinfo['bmHeight']),
                                                  bmpstr, 'raw', 'BGRA', 0, 1)

                    # Overlay the cursor onto the screenshot
                    paste_x = cursor_x - hotspot_x
                    paste_y = cursor_y - hotspot_y

                    img.paste(cursor_img, (paste_x, paste_y), cursor_img)
                else:
                    print("Cursor image is not 32-bit. Skipping overlay.")

                # Cleanup
                win32gui.DeleteObject(icon_info[3])  # mask
                win32gui.DeleteObject(icon_info[4])  # color
                hdc_mem.DeleteDC()
                win32gui.ReleaseDC(0, hdc_screen)

            except Exception as e:
                print(f"Error capturing cursor: {e}")

        # Save final image
        img.save(save_path, optimize=True)    

    def preview_screenshots(self):
        if not self.screenshots:
            self.status_label.configure(text="No screenshots to preview", text_color="orange")
            return

        preview_window = ctk.CTkToplevel(self.root)
        preview_window.transient(self.root)    # Keep it on top of main window
        preview_window.lift()                  # Bring it to front
        preview_window.focus_force()          # Grab keyboard focus
        preview_window.title("Screenshot Viewer")
        
        # Get screen width and height
        preview_window.update_idletasks()  # Ensure dimensions are correct before placing

        # Desired window size (same as your geometry)
        win_width = 1000
        win_height = 700
        
        screen_width = preview_window.winfo_screenwidth()
        screen_height = preview_window.winfo_screenheight()

        # Calculate position
        x = int(((screen_width // 2) - (win_width // 2)) * 2)
        y = int(((screen_height // 2) - (win_height // 2)) * 2)

        # Set centered position
        preview_window.geometry(f"{win_width}x{win_height}+{x}+{y}")

        current_index = [0]  # Mutable index for inner scope

        # Display image area
        img_label = ctk.CTkLabel(preview_window, text="")
        img_label.pack(expand=True, fill="both", padx=10, pady=10)

        def show_image(index):
            if 0 <= index < len(self.screenshots):
                try:
                    img = Image.open(self.screenshots[index])
                    img_resized = img.resize((900, 600), Image.LANCZOS)
                    img_ctk = CTkImage(light_image=img_resized, size=img_resized.size)
                    img_label.configure(image=img_ctk)
                    img_label.image = img_ctk
                    img.close()
                except Exception as e:
                    self.status_label.configure(text=f"Error loading image: {str(e)}", text_color="red")

        def next_image():
            if current_index[0] < len(self.screenshots) - 1:
                current_index[0] += 1
                show_image(current_index[0])

        def prev_image():
            if current_index[0] > 0:
                current_index[0] -= 1
                show_image(current_index[0])

        def delete_current():
            if self.screenshots:
                to_delete = self.screenshots[current_index[0]]
                try:
                    os.remove(to_delete)
                    self.screenshots.remove(to_delete)
                    self.status_label.configure(text=f"Deleted: {Path(to_delete).name}", text_color="orange")
                except Exception as e:
                    self.status_label.configure(text=f"Error deleting: {str(e)}", text_color="red")

            if not self.screenshots:
                preview_window.destroy()
                return

            if current_index[0] >= len(self.screenshots):
                current_index[0] = len(self.screenshots) - 1
            show_image(current_index[0])

        nav_frame = ctk.CTkFrame(preview_window)
        nav_frame.pack(pady=10)

        prev_button = ctk.CTkButton(nav_frame, text="Previous", command=prev_image)
        prev_button.grid(row=0, column=0, padx=10)

        delete_button = ctk.CTkButton(nav_frame, text="Delete", command=delete_current)
        delete_button.grid(row=0, column=1, padx=10)

        next_button = ctk.CTkButton(nav_frame, text="Next", command=next_image)
        next_button.grid(row=0, column=2, padx=10)

        show_image(current_index[0])

    def export_screenshots(self):
        screenshots_to_export = self.selected_screenshots or self.screenshots
        if not screenshots_to_export:
            self.status_label.configure(text="No screenshots to export!", text_color="red")
            return

        save_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("Word Documents", "*.docx")],
            title="Export Screenshots"
        )

        if not save_path:
            return

        try:
            if save_path.endswith(".pdf"):
                from fpdf import FPDF
                from fpdf.enums import XPos, YPos

                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)

                for i, screenshot_path in enumerate(screenshots_to_export, start=1):
                    pdf.add_page()
                    pdf.set_font("Helvetica", size=12)
                    pdf.cell(200, 10, text=f"Step {i}", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
                    pdf.image(str(screenshot_path), x=10, y=20, w=190)

                pdf.output(save_path)

            elif save_path.endswith(".docx"):
                doc = Document()
                doc.add_heading("Screenshot Documentation", level=1)

                for i, screenshot_path in enumerate(screenshots_to_export, start=1):
                    doc.add_paragraph(f"Step {i}")
                    doc.add_picture(str(screenshot_path), width=Inches(6))
                    doc.add_paragraph("")

                doc.save(save_path)

            # Cleanup
            for screenshot_path in screenshots_to_export:
                try:
                    os.remove(screenshot_path)
                except Exception as e:
                    self.status_label.configure(text=f"Error deleting: {screenshot_path.name}", text_color="red")

            self.screenshots.clear()
            self.selected_screenshots.clear()
            self.status_label.configure(text="Exported and Screenshots Deleted!", text_color="green")

        except Exception as e:
            self.status_label.configure(text=f"Export failed: {str(e)}", text_color="red")
            
    def on_closing(self):
        # Delete any remaining screenshots
        for screenshot_path in self.screenshots:
            try:
                os.remove(screenshot_path)
            except Exception as e:
                print(f"Failed to delete {screenshot_path}: {e}")
        self.screenshots.clear()

        # Destroy the main window
        self.root.destroy()
        
    def load_gif(self, gif_path):
        try:
            gif = Image.open(gif_path)
            self.gif_frames = [ImageTk.PhotoImage(frame.copy().convert("RGBA")) for frame in ImageSequence.Iterator(gif)]
            self.gif_index = 0
        except Exception as e:
            print(f"Error loading GIF: {e}")

# Run the app
if __name__ == "__main__":
    import time
    root = ctk.CTk()
    app = ScreenshotApp(root)
    root.mainloop()
